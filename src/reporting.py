import json
from pathlib import Path
from docx import Document
from logger import append_log
from config import EVIDENCE_DIR, TEAM, REG

def generate_json_summary():
    p = Path(EVIDENCE_DIR)
    summary = {"team": TEAM, "reg": REG, "items": []}

    for f in p.glob("*"):
        if f.is_file():
            summary["items"].append({
                "file": f.name,
                "size": f.stat().st_size
            })

    out = p / f"report_manifest_{REG}.json"
    out.write_text(json.dumps(summary, indent=2))

    append_log({
        "module": "reporting",
        "action": "generate_manifest",
        "out": str(out)
    })

    return str(out)

def generate_docx_report():
    p = Path(EVIDENCE_DIR)
    doc = Document()

    doc.add_heading(f"{TEAM} Toolkit Report", 0)
    doc.add_paragraph(f"Team: {TEAM}")
    doc.add_paragraph(f"Reg token: {REG}")
    doc.add_paragraph(
        f"Generated: {__import__('datetime').datetime.utcnow().isoformat()}Z"
    )

    doc.add_heading("Evidence files", level=1)
    for f in sorted(p.glob("*")):
        if f.is_file():
            doc.add_paragraph(f.name)

    # Save inside evidence folder (not root)
    out = p / f"report_{TEAM}.docx"
    doc.save(out)

    append_log({
        "module": "reporting",
        "action": "docx_generated",
        "out": str(out)
    })

    return str(out)
